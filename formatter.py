import re
from dataclasses import dataclass
from io import BytesIO
from typing import Dict, List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONT_CN_SONG = "\u5b8b\u4f53"
FONT_CN_HEI = "\u9ed1\u4f53"
FONT_CN_KAI = "\u6977\u4f53_GB2312"
FONT_EN = "Times New Roman"

SENTENCE_ENDINGS: Sequence[str] = (
    "\u3002",
    "\uff01",
    "\uff1f",
    "\uff1b",
    "\uff1a",
    ".",
    "!",
    "?",
    ";",
    ":",
)


@dataclass
class Block:
    block_type: str
    text: str
    level: int = 0


def _normalize_text(raw_text: str) -> str:
    if not raw_text:
        return ""
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines).strip()


def _heading_level(paragraph: str) -> Optional[int]:
    p = paragraph.strip()
    if not p:
        return None

    special_level1 = {
        "\u6458\u8981",  # 摘要
        "ABSTRACT",
        "\u76ee\u5f55",  # 目录
        "\u53c2\u8003\u6587\u732e",  # 参考文献
        "\u7ed3\u675f\u8bed",  # 结束语
        "\u81f4\u8c22",  # 致谢
    }
    if p.upper() in special_level1:
        return 1

    if re.match(r"^第[0-9\u4e00-\u9fa5]+[章节部分篇]\s*", p):
        return 1
    if re.match(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+、\s*", p):
        return 1
    if re.match(r"^\d+\.\d+\.\d+\s*", p):
        return 3
    if re.match(r"^\d+\.\d+\s*", p):
        return 2
    if re.match(r"^\d+[、.]\s+", p):
        return 1
    if re.match(r"^[（(][0-9\u4e00-\u9fa5]+[）)]\s*", p):
        return 2
    if re.match(r"^\d+、\s*", p):
        return 3

    return None


def _split_paragraphs(text: str) -> List[str]:
    if not text:
        return []

    lines = text.split("\n")
    paragraphs: List[str] = []
    buffer = ""

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if buffer:
                paragraphs.append(buffer)
                buffer = ""
            continue

        if _heading_level(line) is not None:
            if buffer:
                paragraphs.append(buffer)
                buffer = ""
            paragraphs.append(line)
            continue

        if not buffer:
            buffer = line
            continue

        if buffer.endswith(SENTENCE_ENDINGS):
            paragraphs.append(buffer)
            buffer = line
        else:
            buffer += line

    if buffer:
        paragraphs.append(buffer)

    return paragraphs


def _is_title(paragraph: str) -> bool:
    p = paragraph.strip()
    if not p:
        return False
    if len(p) > 45:
        return False
    if _heading_level(p) is not None:
        return False
    if p.endswith(SENTENCE_ENDINGS):
        return False
    return True


def analyze_text(raw_text: str) -> Dict:
    text = _normalize_text(raw_text)
    paragraphs = _split_paragraphs(text)

    if not paragraphs:
        return {
            "title": "",
            "blocks": [],
            "stats": {"paragraph_count": 0, "heading_count": 0},
        }

    title = ""
    start_index = 0
    if _is_title(paragraphs[0]):
        title = paragraphs[0]
        start_index = 1

    blocks: List[Block] = []
    heading_count = 0

    for paragraph in paragraphs[start_index:]:
        level = _heading_level(paragraph)
        if level is not None:
            blocks.append(Block(block_type="heading", text=paragraph, level=level))
            heading_count += 1
        else:
            blocks.append(Block(block_type="paragraph", text=paragraph))

    return {
        "title": title,
        "blocks": [block.__dict__ for block in blocks],
        "stats": {
            "paragraph_count": len(paragraphs),
            "heading_count": heading_count,
        },
    }


def render_preview(structured: Dict) -> str:
    lines: List[str] = []
    title = structured.get("title", "")
    if title:
        lines.append(title)
        lines.append("")

    for block in structured.get("blocks", []):
        if block["block_type"] == "heading":
            lines.append(block["text"])
        else:
            lines.append("\u3000\u3000" + block["text"])
        lines.append("")

    return "\n".join(lines).strip()


def _set_char_spacing(run, spacing: int) -> None:
    if spacing is None:
        return
    r_pr = run._element.get_or_add_rPr()
    spacing_node = r_pr.find(qn("w:spacing"))
    if spacing_node is None:
        spacing_node = OxmlElement("w:spacing")
        r_pr.append(spacing_node)
    spacing_node.set(qn("w:val"), str(spacing))


def _apply_run_font(
    run,
    cn_font: str,
    size_pt: float,
    bold: bool = False,
    align_ascii_font: str = FONT_EN,
    char_spacing: int = 0,
) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = align_ascii_font

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:ascii"), align_ascii_font)
    r_fonts.set(qn("w:hAnsi"), align_ascii_font)

    _set_char_spacing(run, char_spacing)


def _set_base_document_style(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal_font = normal.font
    normal_font.name = FONT_EN
    normal_font.size = Pt(12)

    r_pr = normal._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), FONT_CN_SONG)
    r_fonts.set(qn("w:ascii"), FONT_EN)
    r_fonts.set(qn("w:hAnsi"), FONT_EN)

    normal_fmt = normal.paragraph_format
    normal_fmt.line_spacing = 1.5
    normal_fmt.space_before = Pt(0)
    normal_fmt.space_after = Pt(0)
    normal_fmt.first_line_indent = Pt(24)


def _add_title(document: Document, title: str) -> None:
    if not title:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)

    run = paragraph.add_run(title)
    # 小二号黑体，居中，上下各空一行（近似实现）
    _apply_run_font(run, cn_font=FONT_CN_HEI, size_pt=18, bold=True)


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    run = paragraph.add_run(text)

    if level == 1:
        # 一级标题：三号黑体居中，上下各空一行
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(12)
        _apply_run_font(run, cn_font=FONT_CN_HEI, size_pt=16, bold=True)
    elif level == 2:
        # 二级标题：小四黑体，缩进两个汉字
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(24)
        _apply_run_font(run, cn_font=FONT_CN_HEI, size_pt=12, bold=True)
    else:
        # 三级标题：小四楷体GB2312，缩进两个汉字
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(24)
        _apply_run_font(run, cn_font=FONT_CN_KAI, size_pt=12, bold=False)


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5

    run = paragraph.add_run(text)
    # 正文：中文小四宋体、英文小四Times New Roman
    _apply_run_font(run, cn_font=FONT_CN_SONG, size_pt=12, bold=False)


def build_docx_bytes(structured: Dict) -> bytes:
    document = Document()
    _set_base_document_style(document)
    _add_title(document, structured.get("title", ""))

    for block in structured.get("blocks", []):
        if block["block_type"] == "heading":
            _add_heading(document, block["text"], block["level"])
        else:
            _add_body_paragraph(document, block["text"])

    output = BytesIO()
    document.save(output)
    return output.getvalue()
